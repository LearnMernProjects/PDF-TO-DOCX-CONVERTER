from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns


# -------------------------------------------------
# TABLE BORDER HELPER
# -------------------------------------------------
def set_table_borders(table):
    """
    Forces visible borders on a Word table.
    Without this, table lines may not appear.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblBorders = OxmlElement("w:tblBorders")

    for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge = OxmlElement(f"w:{border}")
        edge.set(ns.qn("w:val"), "single")
        edge.set(ns.qn("w:sz"), "12")      # border thickness
        edge.set(ns.qn("w:space"), "0")
        edge.set(ns.qn("w:color"), "000000")
        tblBorders.append(edge)

    tblPr.append(tblBorders)


# -------------------------------------------------
# COLUMN WIDTH HELPER
# -------------------------------------------------
def set_column_widths(table, widths):
    """
    widths: list of Inches, one per column
    Locks column widths to match PDF layout.
    """
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


# -------------------------------------------------
# CELL ALIGNMENT HELPER
# -------------------------------------------------
def align_cell(cell, alignment: str):
    """
    Applies inferred alignment to a Word table cell.
    """
    for paragraph in cell.paragraphs:
        if alignment == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


# -------------------------------------------------
# TABLE CREATION
# -------------------------------------------------
def add_form_table(document, table_data):
    """
    Adds a bordered, aligned form-style table to the document.
    """
    # Skip if no table data
    if not table_data or len(table_data) == 0:
        return
    
    rows = len(table_data)
    cols = len(table_data[0])

    table = document.add_table(rows=rows, cols=cols)

    # Apply visible grid style
    table.style = "Table Grid"

    # Force borders explicitly
    set_table_borders(table)

    # Lock column widths for legal form (3-column layout)
    if cols == 3:
        set_column_widths(
            table,
            [Inches(0.6), Inches(2.5), Inches(3.0)]
        )

    # Fill table cells
    for r_idx, row in enumerate(table_data):
        for c_idx, cell_value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_value or ""

            # Default alignment for form tables
            align_cell(cell, "left")

    return table


# -------------------------------------------------
# MAIN WORD DOCUMENT BUILDER
# -------------------------------------------------
def build_word_document(structured_doc: dict, output_path: str, aligned_lines=None):

    """
    Builds a Word document from structured text and tables.

    Input:
    - structured_doc: output of STEP 2 (text classification)
    - output_path: path to save .docx file
    """

    document = Document()

    # -------------------------
    # TITLE
    # -------------------------
    if structured_doc.get("title"):
        title_para = document.add_paragraph()
        title_run = title_para.add_run(structured_doc["title"])
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")

    # -------------------------
    # SECTIONS (TEXT)
    # -------------------------
    line_counter = 0  # Track line count for alignment
    for heading, lines in structured_doc["sections"].items():
        heading_para = document.add_paragraph()
        heading_run = heading_para.add_run(heading)
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        
        for idx, line in enumerate(lines):
            para = document.add_paragraph(line)
            para.paragraph_format.space_after = Pt(6)

            # Apply semantic alignment if available
            if aligned_lines and line_counter < len(aligned_lines):
                align = aligned_lines[line_counter].get("alignment", "left")

                if align == "right":
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == "center":
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            line_counter += 1




    # -------------------------
    # EXAMPLE: FORM TABLE INSERTION
    # (replace table_data with extracted table later)
    # -------------------------
    table_data = [
    
    ]

    add_form_table(document, table_data)

    # -------------------------
    # SAVE DOCUMENT
    # -------------------------
    document.save(output_path)
